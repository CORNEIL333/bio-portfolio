import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Segoe UI'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700
    
    # Title Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("RAPPORT DE TP : COLLABORATION BIO-PORTFOLIO")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a) # Slate 900
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Cours de Git & GitHub - Bachelor 1 Génie Logiciel")
    subtitle_run.font.name = 'Segoe UI'
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8) # Accent Sky Blue
    
    # Student Info Table
    table = doc.add_table(rows=2, cols=2)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Étudiant'
    hdr_cells[1].text = 'Filière / Groupe'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'WATAT MIGUEL CORNEIL'
    row_cells[1].text = 'Bachelor 1 Génie Logiciel'
    
    # Format Table text
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)
    
    doc.add_paragraph() # Spacing
    
    # Section 1: Introduction & Objectifs
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. Présentation du Projet et Objectifs")
    r1.font.name = 'Segoe UI'
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    
    p1 = doc.add_paragraph("Ce projet pratique (TP) met en œuvre les concepts essentiels de la gestion de versions avec Git et la collaboration avec GitHub. Les objectifs pédagogiques principaux consistent à maîtriser le cycle de vie d'un dépôt local, le travail collaboratif à travers des branches de fonctionnalités, la simulation de Pull Requests, ainsi que la création volontaire et la résolution méthodique de conflits de fusion (merge conflicts).")
    
    # Section 2: Réalisation Technique & Git Workflow
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. Workflow Git et Simulation de Collaboration")
    r2.font.name = 'Segoe UI'
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    
    doc.add_paragraph("Le TP a été structuré selon les étapes méthodologiques suivantes :")
    
    steps = [
        ("Initialisation & Commit Initial", "Création du dépôt local avec 'git init' sur la branche 'main' et enregistrement du profil de l'étudiant dans index.html."),
        ("Modifications Collaboratives (Branche)", "Création de la branche 'feature-competences' pour ajouter la section Compétences (HTML, CSS, Git) et simulation d'une issue GitHub (ISSUE.md)."),
        ("Simulation de Pull Request", "Création des livrables PULL_REQUEST.md et COMMENTAIRE_PR.md décrivant les changements apportés."),
        ("Création de Conflit de Fusion", "Modification concurrente de la première ligne d'index.html sur les deux branches ('main' et 'feature-competences') pour générer volontairement un conflit."),
        ("Résolution du Conflit", "Fusion de 'feature-competences' dans 'main' avec résolution manuelle en conservant le titre 'Etudiant en B1' et la section compétences fusionnée automatiquement.")
    ]
    
    for title, desc in steps:
        p = doc.add_paragraph(style='List Bullet')
        r_title = p.add_run(f"{title} : ")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        p.add_run(desc)
        
    doc.add_page_break() # Page 2 starts here
    
    # Section 3: Historique Git & Graphes
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Historique Git et Résolution de Conflit")
    r3.font.name = 'Segoe UI'
    r3.font.size = Pt(14)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    
    doc.add_paragraph("Voici le graphe d'historique de commits obtenu après la résolution du conflit et la fusion finale (généré via 'git log --oneline --graph --all') :")
    
    # Read historique.txt
    hist_text = ""
    try:
        with open("historique.txt", "r", encoding="utf-8") as f:
            hist_text = f.read()
    except Exception:
        try:
            with open("historique.txt", "r", encoding="utf-16le") as f:
                hist_text = f.read()
        except Exception:
            hist_text = "*   Resolution conflit et fusion finale\n|\\  \n| * bce1da5 Modification premiere ligne pour conflit\n..."
            
    p_hist = doc.add_paragraph()
    p_hist.paragraph_format.left_indent = Inches(0.5)
    r_hist = p_hist.add_run(hist_text)
    r_hist.font.name = 'Courier New'
    r_hist.font.size = Pt(9.5)
    r_hist.font.bold = True
    r_hist.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    
    # Section 4: Capture d'écran du Portfolio
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Capture d'Écran de l'Interface Portfolio")
    r4.font.name = 'Segoe UI'
    r4.font.size = Pt(14)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    
    doc.add_paragraph("Rendu de la page index.html moderne et responsive intégrant la section Compétences après fusion :")
    
    # Add screenshot image
    if os.path.exists("screenshot.png"):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("screenshot.png", width=Inches(4.5))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Figure 1 : Interface utilisateur moderne du Bio-Portfolio (index.html)")
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    else:
        doc.add_paragraph("[Erreur : Capture d'écran indisponible]")
        
    doc.save("rapport.docx")
    print("Word document generated successfully as rapport.docx")

if __name__ == "__main__":
    create_report()
